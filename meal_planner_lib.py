import json
import random
import webbrowser

import requests
from docx import Document

from flask import render_template, request

"""
----------------------------------------------------------------------------
Globals
----------------------------------------------------------------------------
"""

# global to store recipes
new_data = {"hits": []}
selected_data = {"hits": []}


def create_recipe_document():
    total_recipes = []
    just_ingredients = []
    modified_data = [total_recipes, just_ingredients]

    # create new document
    doc = Document()
    # add heading
    doc.add_heading("Saved Recipes")

    for i, recipe_data in enumerate(new_data['hits'], start=1):
        # rename the recipe so they go in ascending order
        new_name = 'recipe ' + str(i)

        # create the new recipe dictionary
        new_recipe = {new_name: recipe_data}

        # add new recipe to list
        total_recipes.append(new_recipe)

        # isolate the ingredients list from the recipe
        new_ingredients = {new_name + ' ingredients': recipe_data['recipe']['ingredientLines']}

        # add isolated ingredients to list of ingredients
        just_ingredients.append(new_ingredients)

    # add each recipe to document
    for recipe_info in modified_data[0]:
        for result_number, recipe_details in recipe_info.items():
            title_paragraph = doc.add_paragraph()
            runner = title_paragraph.add_run(f"Recipe Title: {recipe_details['recipe']['label']}")
            runner.bold = True
            doc.add_paragraph(f"URL: {recipe_details['recipe']['url']}")

            # add ingredients as bulleted list
            doc.add_paragraph(f"Ingredients: ")
            for each_ingredient in recipe_details['recipe']['ingredientLines']:
                ingredient_paragraph = doc.add_paragraph(f"{each_ingredient}")
                ingredient_paragraph.style = 'List Bullet'

            doc.add_paragraph("\n")

    response_filename = 'Recipes.docx'
    doc.save(response_filename)


def create_ingredients_document():
    total_recipes = []
    just_ingredients = []
    modified_data = [total_recipes, just_ingredients]
    for i, recipe_data in enumerate(new_data['hits'], start=1):
        # rename the recipe so they go in ascending order
        new_name = 'recipe ' + str(i)

        # create the new recipe dictionary
        new_recipe = {new_name: recipe_data}

        # add new recipe to list
        total_recipes.append(new_recipe)

        # isolate the ingredients list from the recipe
        new_ingredients = {new_name + ' ingredients': recipe_data['recipe']['ingredientLines']}

        # add isolated ingredients to list of ingredients
        just_ingredients.append(new_ingredients)

    # create new document
    doc = Document()
    # add heading
    doc.add_heading("Ingredients List")

    # add each recipe to document
    for recipe_info in modified_data[0]:
        for result_number, recipe_details in recipe_info.items():
            for recipe_ingredient in recipe_details['recipe']['ingredientLines']:
                doc.add_paragraph(f"{recipe_ingredient}")

    response_filename = 'Ingredients List.docx'
    doc.save(response_filename)


def open_url(url):
    import webbrowser
    webbrowser.open(url)


def argument_handler(*args):
    if not args:
        return "You need to provide at least 1 argument."
    else:
        return list(args)


def remove_str_chars(input_string, num):
    if num >= 0:
        return input_string[:-num]
    else:
        return input_string


def process_food_list(input_list):
    request_string = ""
    for food in input_list:
        request_string = request_string + food + "%2c%20"
    # remove the last %2c%20 from the collated string which is unicode for ", "
    api_string = remove_str_chars(request_string, 6)
    return api_string


def browse_recipes():
    url = "https://www.allrecipes.com/"
    webbrowser.open(url)


def search_recipes(excluded_ingredients_entry_param, ingredients_entry_param, num_recipes_entry_param):
    global selected_data
    global new_data
    excluded_ingredients_str = request.form.get('excluded_ingredients')
    ingredients = request.form.get('ingredients')
    if not ingredients:
        render_template("Please enter the number of recipes.\n")
        return
    try:
        num_recipes = int(request.form.get('num_recipes'))
    except ValueError:
        render_template("Please input a valid number.\n")
        return
    if num_recipes <= 0 or num_recipes > 20:
        render_template("Please enter a valid positive number for the number of recipes (between 1 and 20).\n")
        return

    if ingredients:
        food_list = argument_handler(ingredients)
        formatted_string = process_food_list(food_list)
        response = requests.get("https://api.edamam.com/api/recipes/v2?type=public&q=" + formatted_string +
                                "&app_id=2286dd85&app_key=1cdfcd395ccf99e349b18f54eaa4416f&" + excluded_ingredients_str
                                + "&random=true&field=url&field=label&field=ingredientLines")
        dict_from_json = json.loads(response.text)
        if not dict_from_json["hits"]:
            render_template(f"Your search for {ingredients} found no recipes, please try again.")
            return
    else:
        render_template("Please select at least one ingredient.\n")
        return

    if response.status_code == 200:
        # parse the json
        dict_from_json = json.loads(response.text)
        if not dict_from_json["hits"]:
            render_template(f"Your search for {ingredients} found no recipes, please try again.")
            return
        selected_recipes = random.sample(dict_from_json["hits"], num_recipes)
        selected_data = {
            "hits": selected_recipes
        }
        recipes_data = []
        for i, recipe_data in enumerate(selected_data["hits"], start=1):
            recipe = recipe_data["recipe"]
            recipe_url = recipe["url"]
            recipe_name = recipe["label"]
            ingredients = recipe_data["recipe"]["ingredientLines"]
            recipes_data.append({
                "recipe_name": recipe_name,
                "recipe_url": recipe_url,
                "ingredients": ingredients
            })
        return None, recipes_data

    else:
        render_template(f"API request failed with status code: {response.status_code}")


def save_recipes(output_text):
    saved_recipes = output_text
    if not saved_recipes:
        render_template(f"Please select more than 0 recipes and make sure every selection is valid.\n")
        return
    split_values = saved_recipes.split(',')
    # remove whitespace in list
    current_saved_recipe_list = []
    for x in split_values:
        stripped_x = x.strip()
        current_saved_recipe_list.append(stripped_x)

    # removes non number values, but leaves a single negative sign if found (-)
    integer_list = []
    for x in current_saved_recipe_list:
        stripped_x = x.lstrip('-')
        if stripped_x.isdigit():
            integer_list.append(int(x))

    if int(min(integer_list)) <= 0:
        render_template(f"Please select more than 0 recipes and make sure every selection is valid.\n")
        return

    for val in current_saved_recipe_list:
        try:
            int(val)
        except ValueError:
            render_template("Please input a valid number.\n")
            return
        if int(val) < 0 or int(val) > len(selected_data["hits"]):
            render_template(f"Please make sure your number is greater than 0 and less than total recipes searched.\n")
            return

    output_text.insert("1.0", f"Here is what you selected: {integer_list}\n")
    output_text.insert("2.0", "******* Adding recipes to saved recipes... *******\n\n")

    for idx in range(len(integer_list)):
        new_data["hits"].append(selected_data["hits"][integer_list[idx] - 1])

    return
